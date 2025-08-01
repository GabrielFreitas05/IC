from docx import Document
from docx.shared import Inches
import os
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

def inserir_paragrafo_apos(paragrafo):
    body = paragrafo._parent._element
    p_element = paragrafo._element
    new_p = OxmlElement('w:p')
    body.insert(body.index(p_element) + 1, new_p)
    novo_paragrafo = Paragraph(new_p, paragrafo._parent)
    return novo_paragrafo

def preencher_pop(dados, caminho_modelo, imagens, legendas):
    doc = Document(caminho_modelo)

    for paragrafo in doc.paragraphs:
        texto_original = ''.join(run.text for run in paragrafo.runs)
        for chave, valor in dados.items():
            if chave in ("imagem", "anexos"):
                continue
            placeholder = f"{{{{{chave}}}}}"
            if placeholder in texto_original:
                novo_texto = texto_original.replace(placeholder, valor)
                for run in paragrafo.runs:
                    run.text = ''
                paragrafo.runs[0].text = novo_texto
                texto_original = novo_texto

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    texto_original = ''.join(run.text for run in paragrafo.runs)
                    for chave, valor in dados.items():
                        if chave in ("imagem", "anexos"):
                            continue
                        placeholder = f"{{{{{chave}}}}}"
                        if placeholder in texto_original:
                            novo_texto = texto_original.replace(placeholder, valor)
                            for run in paragrafo.runs:
                                run.text = ''
                            paragrafo.runs[0].text = novo_texto
                            texto_original = novo_texto

    if imagens:
        historico_presente = 'historico_revisoes' in dados and dados['historico_revisoes'].strip()
        for paragrafo in doc.paragraphs:
            if '{{anexos}}' in paragrafo.text:
                paragrafo.text = paragrafo.text.replace('{{anexos}}', '')
                ultimo_paragrafo = paragrafo
                for i, (imagem, legenda) in enumerate(zip(imagens, legendas)):
                    novo_p_img = inserir_paragrafo_apos(ultimo_paragrafo)
                    run_img = novo_p_img.add_run()
                    try:
                        run_img.add_picture(imagem, width=Inches(4))
                    except Exception as e:
                        novo_p_img.add_run(f"[Erro ao inserir imagem: {str(e)}]")
                    ultimo_paragrafo = novo_p_img

                    novo_p_leg = inserir_paragrafo_apos(ultimo_paragrafo)
                    novo_p_leg.add_run(f'Figura {i+1}: {legenda}')
                    ultimo_paragrafo = novo_p_leg

                break

    nome_arquivo = f"{dados['titulo_procedimento']}_POP.docx"
    caminho_temporario = os.path.join("assets", nome_arquivo)
    doc.save(caminho_temporario)
    return caminho_temporario
